import os
import pandas as pd
import numpy as np
from docx import Document

def transform_table_to_scale(df):
    if "Strongly Disagree" in df.columns:
        column_mapping = {
            "Strongly Disagree": 1,
            "Somewhat Disagree": 2,
            "Neutral": 3,
            "Somewhat Agree": 4,
            "Strongly Agree": 5
        }
        transformed_df = pd.DataFrame(columns=["Question", "Scale"])
        for index, row in df.iterrows():
            for col, scale in column_mapping.items():
                if row[col].strip():  # Non-empty cells are considered
                    transformed_df = transformed_df.append({"Question": row.name, "Scale": scale}, ignore_index=True)
                    break
    else:
        transformed_df = pd.DataFrame(columns=["Question", "Scale"])
        for index, row in df.iterrows():
            for col in df.columns:
                if row[col].strip():  # Non-empty cells are considered
                    try:
                        scale_value = int(col.split()[0])
                    except IndexError:
                        scale_value = col  # Use the full column name if it doesn't contain any spaces
                    transformed_df = transformed_df.append({"Question": row.name, "Scale": scale_value}, ignore_index=True)
                    break
    return transformed_df

def get_response_between_markers(text, start_marker, end_marker):
    start = text.find(start_marker) + len(start_marker)
    end = text.find(end_marker)
    return text[start:end].strip()

# Initialize an empty DataFrame to store the final output
final_df = pd.DataFrame()

# List all .docx files in the current directory
filenames = [f for f in os.listdir() if f.endswith('.docx')]

for filename in filenames:
    file_path = filename
    document = Document(file_path)
    
    tables_list = []
    for table in document.tables:
        rows_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            rows_data.append(row_data)
        tables_list.append(pd.DataFrame(rows_data[1:], columns=rows_data[0]))
        
    transformed_tables = [transform_table_to_scale(df) for df in tables_list]
    
    data_row = [filename]
    for transformed_table in transformed_tables:
        scales = transformed_table['Scale'].tolist()
        data_row.extend(scales)
        
    # Capture all text in the document
    full_text = "\n".join([p.text for p in document.paragraphs])
    
    # Capture the responses for questions 15, 16, and 17 based on the markers
    response_15 = get_response_between_markers(full_text, "What was the most rewarding experience for you during the REU project?", "16. What was the most frustrating experience for you during the REU project?")
    response_16 = get_response_between_markers(full_text, "What was the most frustrating experience for you during the REU project?", "17. Any additional comments and suggestions are welcome below:")
    response_17 = get_response_between_markers(full_text, "Any additional comments and suggestions are welcome below:", "Thank you")
    
    # Append the responses for questions 15, 16, and 17 to data_row
    data_row.extend([response_15, response_16, response_17])
    
    temp_df = pd.DataFrame([data_row], columns=['Filename'] + [f'Scale_Q{i+1}' for i in range(len(data_row) - 1)])
    final_df = final_df.append(temp_df, ignore_index=True)

# Save the final DataFrame to an Excel file
final_df.to_excel("Reformatted_Scales_Multiple_Files.xlsx", index=False)
